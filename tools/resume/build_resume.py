from __future__ import annotations

import argparse
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


# Visual authority: William's July 2025 CV PDF, distilled in artifact.md.
# Intentional changes: one page, Applied AI focus, city-level address, 10 pt body.
FONT = "Arial"
INK = RGBColor(0x00, 0x00, 0x00)
LINK_BLUE = RGBColor(0x05, 0x61, 0xC9)
BODY_SIZE = 10.5
CONTENT_WIDTH_IN = 7.40
BULLET_TEXT_INDENT_IN = 0.36
BULLET_HANGING_IN = 0.16


def set_style_font(style, *, size: float, bold: bool = False, italic: bool = False):
    style.font.name = FONT
    fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    fonts.set(qn("w:eastAsia"), FONT)
    style.font.size = Pt(size)
    style.font.color.rgb = INK
    style.font.bold = bold
    style.font.italic = italic


def set_run_font(run, *, size: float = BODY_SIZE, bold=None, italic=None, color=INK):
    run.font.name = FONT
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    fonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_keep_with_next(paragraph, value: bool = True):
    paragraph.paragraph_format.keep_with_next = value


def set_bottom_rule(paragraph):
    properties = paragraph._p.get_or_add_pPr()
    borders = properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        properties.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "000000")
    borders.append(bottom)


def add_hyperlink(paragraph, text: str, url: str, *, size: float = 10.0):
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    fonts.set(qn("w:eastAsia"), FONT)
    properties.append(fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0561C9")
    properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    word_size = str(round(size * 2))
    for tag in ("w:sz", "w:szCs"):
        size_node = OxmlElement(tag)
        size_node.set(qn("w:val"), word_size)
        properties.append(size_node)
    run.append(properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def create_bullet_numbering(document: Document) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    number_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    number_id = max(number_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    number_format = OxmlElement("w:numFmt")
    number_format.set(qn("w:val"), "bullet")
    level.append(number_format)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "\u2022")
    level.append(level_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)

    paragraph_properties = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "518")
    tabs.append(tab)
    paragraph_properties.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "518")
    indent.set(qn("w:hanging"), "230")
    paragraph_properties.append(indent)
    level.append(paragraph_properties)

    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    run_properties.append(fonts)
    level.append(run_properties)
    abstract.append(level)
    numbering.append(abstract)

    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(number_id))
    abstract_reference = OxmlElement("w:abstractNumId")
    abstract_reference.set(qn("w:val"), str(abstract_id))
    number.append(abstract_reference)
    numbering.append(number)
    return number_id


def apply_numbering(paragraph, number_id: int):
    properties = paragraph._p.get_or_add_pPr()
    number_properties = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number_properties.append(level)
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(number_id))
    number_properties.append(number)
    properties.append(number_properties)


def add_bullet(document: Document, number_id: int, text: str):
    paragraph = document.add_paragraph(style="Resume Bullet")
    apply_numbering(paragraph, number_id)
    paragraph.paragraph_format.left_indent = Inches(BULLET_TEXT_INDENT_IN)
    paragraph.paragraph_format.first_line_indent = Inches(-BULLET_HANGING_IN)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(BULLET_TEXT_INDENT_IN))
    set_run_font(paragraph.add_run(text), size=BODY_SIZE)
    return paragraph


def add_certification_bullet(document: Document, number_id: int, text: str):
    paragraph = document.add_paragraph(style="Resume Certification")
    apply_numbering(paragraph, number_id)
    paragraph.paragraph_format.left_indent = Inches(0.20)
    paragraph.paragraph_format.first_line_indent = Inches(-0.10)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(0.20))
    set_run_font(paragraph.add_run(text), size=8.5)
    return paragraph


def add_section(document: Document, title: str, *, compact: bool = False):
    paragraph = document.add_paragraph(title.upper(), style="Heading 1")
    for run in paragraph.runs:
        set_run_font(run, size=11.0, bold=True)
    if compact:
        paragraph.paragraph_format.space_before = Pt(7.5)
        paragraph.paragraph_format.space_after = Pt(3.5)
    set_bottom_rule(paragraph)
    set_keep_with_next(paragraph)
    return paragraph


def add_skill(document: Document, label: str, value: str):
    paragraph = document.add_paragraph(style="Resume Skill")
    set_run_font(paragraph.add_run(f"{label}: "), size=BODY_SIZE, bold=True)
    set_run_font(paragraph.add_run(value), size=BODY_SIZE)
    return paragraph


def add_entry_title(
    document: Document,
    title: str,
    links: list[tuple[str, str]] | None = None,
    date: str | None = None,
    *,
    first_in_section: bool = False,
):
    paragraph = document.add_paragraph(style="Resume Entry")
    if first_in_section:
        paragraph.paragraph_format.space_before = Pt(1.5)
    set_run_font(paragraph.add_run(title), size=BODY_SIZE, bold=True)
    for label, url in links or []:
        set_run_font(paragraph.add_run(" | "), size=BODY_SIZE)
        add_hyperlink(paragraph, label, url, size=BODY_SIZE)
    if date:
        paragraph.add_run("\t")
        set_run_font(paragraph.add_run(date), size=BODY_SIZE, italic=True)
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            Inches(CONTENT_WIDTH_IN), WD_TAB_ALIGNMENT.RIGHT
        )
    set_keep_with_next(paragraph)
    return paragraph


def add_meta_line(document: Document, details: str, date: str):
    paragraph = document.add_paragraph(style="Resume Meta")
    if details:
        set_run_font(paragraph.add_run(details), size=BODY_SIZE, italic=True)
        paragraph.add_run("\t")
    set_run_font(paragraph.add_run(date), size=BODY_SIZE, italic=True)
    if not details:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(CONTENT_WIDTH_IN), WD_TAB_ALIGNMENT.RIGHT
    )
    set_keep_with_next(paragraph)
    return paragraph


def configure_styles(document: Document, *, compact: bool = False):
    styles = document.styles
    normal = styles["Normal"]
    set_style_font(normal, size=BODY_SIZE)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    heading = styles["Heading 1"]
    set_style_font(heading, size=11.0, bold=True)
    heading.paragraph_format.space_before = Pt(7.5 if compact else 9.0)
    heading.paragraph_format.space_after = Pt(3.5)
    heading.paragraph_format.line_spacing = 1.0
    heading.paragraph_format.keep_with_next = True

    specs = {
        "Resume Name": (20.0, True, False, 0, 0, 1.0),
        "Resume Role": (10.5, True, False, 0, 0.5, 1.0),
        "Resume Contact": (10.0, False, False, 0, 0, 1.0),
        "Resume Contact Last": (10.0, False, False, 0, 4.0 if compact else 5.0, 1.0),
        "Resume Summary": (BODY_SIZE, False, False, 0, 7.5 if compact else 11.0, 1.08),
        "Resume Skill": (BODY_SIZE, False, False, 0, 1.5 if compact else 2.25, 1.08),
        "Resume Entry": (BODY_SIZE, False, False, 5.5 if compact else 6.5, 0.8, 1.0),
        "Resume Meta": (BODY_SIZE, False, True, 0, 1.3 if compact else 2.0, 1.0),
        "Resume Bullet": (BODY_SIZE, False, False, 0, 2.0 if compact else 3.0, 1.08),
        "Resume Detail": (BODY_SIZE, False, False, 0, 1.4 if compact else 2.25, 1.08),
        "Resume Certification": (8.5, False, False, 0, 0.45, 1.0),
    }
    for name, (size, bold, italic, before, after, spacing) in specs.items():
        style = styles[name] if name in styles else styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = normal
        set_style_font(style, size=size, bold=bold, italic=italic)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = spacing


def build_resume(output_path: Path, *, include_certifications: bool = False):
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.35 if include_certifications else 0.48)
    section.bottom_margin = Inches(0.20 if include_certifications else 0.42)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.20)
    section.footer_distance = Inches(0.20)

    configure_styles(document, compact=include_certifications)
    bullet_number_id = create_bullet_numbering(document)

    properties = document.core_properties
    properties.title = "William Lo Channiko - Applied AI Engineer Resume"
    properties.subject = "Applied AI Engineer resume"
    properties.author = "William Lo Channiko"
    properties.comments = "One-page Applied AI Engineer resume."
    properties.language = "en-US"
    properties.last_modified_by = "William Lo Channiko"
    properties.created = datetime.now(timezone.utc)
    properties.modified = datetime.now(timezone.utc)
    properties.revision = 1
    properties.keywords = (
        "Applied AI Engineer, AI Engineer, Python, FastAPI, TypeScript, React, "
        "Next.js, OpenAI API, RAG, LangGraph, PostgreSQL, human-in-the-loop"
    )

    name = document.add_paragraph(style="Resume Name")
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(name.add_run("William Lo Channiko"), size=20.0, bold=True)

    role = document.add_paragraph(style="Resume Role")
    role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(role.add_run("APPLIED AI ENGINEER"), size=10.5, bold=True)

    contact = document.add_paragraph(style="Resume Contact")
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(contact.add_run("Jakarta, Indonesia | +62 813 4852 0623 | "), size=10.0)
    add_hyperlink(contact, "williamlochanniko4@gmail.com", "mailto:williamlochanniko4@gmail.com")

    profile_links = document.add_paragraph(style="Resume Contact Last")
    profile_links.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_hyperlink(
        profile_links,
        "linkedin.com/in/william-lo-channiko",
        "https://www.linkedin.com/in/william-lo-channiko/",
    )
    set_run_font(profile_links.add_run(" | "), size=10.0)
    add_hyperlink(profile_links, "github.com/williamlo90", "https://github.com/williamlo90")
    set_run_font(profile_links.add_run(" | "), size=10.0)
    add_hyperlink(
        profile_links,
        "william-lo-channiko-portfolio.pages.dev",
        "https://william-lo-channiko-portfolio.pages.dev/",
    )

    summary = document.add_paragraph(style="Resume Summary")
    set_run_font(
        summary.add_run(
            "Applied AI engineer with 1 year of full-time software engineering internship experience "
            "across two terms. "
            "Built end-to-end Python/FastAPI and TypeScript/React systems spanning RAG, structured "
            "extraction, human-in-the-loop workflows, offline evaluation, and IndoBERT fine-tuning."
        ),
        size=BODY_SIZE,
    )

    add_section(document, "Technical Skills")
    add_skill(
        document,
        "Core engineering",
        "Python, TypeScript, SQL; FastAPI, Next.js, React; PostgreSQL, REST APIs, Docker, Git, GitHub Actions, CI/CD",
    )
    add_skill(
        document,
        "Applied AI",
        "OpenAI API, LangGraph, RAG, embeddings, pgvector, structured outputs, Mistral OCR, offline evaluation, human-in-the-loop workflows",
    )
    add_skill(
        document,
        "ML/NLP and testing",
        "PyTorch, Hugging Face Transformers, scikit-learn, IndoBERT, Optuna; Pytest, Vitest, Playwright; Kotlin, PHP/CodeIgniter, SQLite",
    )

    add_section(document, "Selected AI Projects")
    add_entry_title(
        document,
        "Case Resolution Copilot",
        [("GitHub Repository", "https://github.com/williamlo90/case-resolution-copilot")],
        date="Jul 2026-Present",
        first_in_section=True,
    )
    add_bullet(
        document,
        bullet_number_id,
        "Built a policy-governed case-resolution workspace with FastAPI, PostgreSQL/pgvector, LangGraph, OpenAI, and controlled Gmail drafts; kept evidence, approvals, and action authority server-controlled. LangChain Core handled bounded formatting; CrewAI and AutoGen stayed isolated one-case comparisons.",
    )
    add_bullet(
        document,
        bullet_number_id,
        "Reduced median workflow time by 84% (582s to 95s) while completing 3/3 safe workflows versus 0/3 manually in a matched synthetic benchmark; validated PostgreSQL persistence, provider-failure handling, and frontend recovery for controlled-pilot readiness.",
    )

    add_entry_title(
        document,
        "Invoice Review",
        [("GitHub Repository", "https://github.com/williamlo90/ai-document-ops-system")],
        date="Jul-Aug 2026",
    )
    add_bullet(
        document,
        bullet_number_id,
        "Built an AI-powered invoice-to-ERP workflow with React, FastAPI, Mistral OCR, OpenAI structured outputs, and ERPNext; added deterministic validation, human review, audit history, and approval-gated idempotent draft delivery.",
    )
    add_bullet(
        document,
        bullet_number_id,
        "Cut median invoice-to-ERP draft time by 68% (153s to 49s) while achieving the expected result in 10/10 cases versus 9/10 through direct entry; retained 98.75% exact field match and 100% validation and blocker match on a sealed synthetic holdout.",
    )

    add_section(document, "Professional Experience")
    add_entry_title(document, "PT Dover Chemical", first_in_section=True)
    add_meta_line(
        document,
        "Software Engineer Intern (Mobile and Full-Stack) | Jakarta, Indonesia",
        "Jan 2025-Jan 2026",
    )
    add_bullet(
        document,
        bullet_number_id,
        "Built offline-first Draft Sales Order workflows for the CRM Dover Chemical Android app, enabling field-sales users to create and retain orders without connectivity; integrated 38 HTTP API operations, 15 Room entities/DAOs, and WorkManager synchronization for records, attachments, retries, and cleanup.",
    )
    add_bullet(
        document,
        bullet_number_id,
        "Built a Customer Management System for a 6,656-record customer master dataset across 15 data domains, including CRUD/search, validated Excel imports, reporting, audit logs, and role-gated approval, rejection, and reopen paths.",
    )

    add_section(document, "Education")
    education = document.add_paragraph(style="Resume Entry")
    education.paragraph_format.space_before = Pt(1.5)
    set_run_font(
        education.add_run("Multimedia Nusantara University"),
        size=BODY_SIZE,
        bold=True,
    )
    set_run_font(education.add_run(", Tangerang, Indonesia"), size=BODY_SIZE)
    education.add_run("\t")
    set_run_font(education.add_run("2022-2026"), size=BODY_SIZE)
    education.paragraph_format.tab_stops.add_tab_stop(
        Inches(CONTENT_WIDTH_IN), WD_TAB_ALIGNMENT.RIGHT
    )
    set_keep_with_next(education)

    degree = document.add_paragraph(style="Resume Detail")
    set_run_font(
        degree.add_run("Bachelor's Degree in Informatics (S.Kom.)"),
        size=BODY_SIZE,
        italic=True,
    )
    set_run_font(
        degree.add_run(" | Graduated with Distinction | GPA: 3.82/4.00"),
        size=BODY_SIZE,
    )
    set_keep_with_next(degree)

    thesis = document.add_paragraph(style="Resume Detail")
    set_run_font(
        thesis.add_run(
            "Thesis: Fine-tuned two IndoBERT classifiers on 1,822 Indonesian social-media opinions; "
            "achieved test macro-F1 of 0.981 for sentiment and 0.870 for aspect classification, with "
            "inter-annotator Cohen's kappa of 0.954 and 0.900 on a 200-opinion reliability sample."
        ),
        size=BODY_SIZE,
    )
    set_keep_with_next(thesis)

    honors = document.add_paragraph(style="Resume Detail")
    set_run_font(
        honors.add_run(
            "Honors: Academic Achievement Scholarship, 2024"
        ),
        size=BODY_SIZE,
    )

    if include_certifications:
        add_section(document, "Certifications", compact=True)
        for certification in (
            "Machine Learning Specialization - DeepLearning.AI, Apr 2025",
            "Machine Learning A-Z: AI, Python & R - Udemy, Mar 2025",
            "Google Data Analytics Professional Certificate, Mar 2025",
            "SQL for Data Science - University of California, Nov 2024",
            "Data Wrangling, Analysis and A/B Testing with SQL - University of California, Nov 2024",
        ):
            add_certification_bullet(document, bullet_number_id, certification)

    for paragraph in document.paragraphs:
        paragraph.paragraph_format.widow_control = True

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--include-certifications", action="store_true")
    arguments = parser.parse_args()
    build_resume(
        arguments.output.resolve(),
        include_certifications=arguments.include_certifications,
    )


if __name__ == "__main__":
    main()
